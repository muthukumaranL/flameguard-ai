"""Render the report block model to Markdown, DOCX and PDF."""
from __future__ import annotations

from pathlib import Path
from typing import Any

from src.utils import setup_logging

log = setup_logging("flameguard.render")

Block = tuple[str, Any]

ACCENT = "#B03A2E"


# ------------------------------------------------------------------ markdown
def render_markdown(blocks: list[Block], out_path: Path) -> None:
    lines: list[str] = []
    for kind, payload in blocks:
        if kind == "cover":
            lines += [f"# {payload['title']}", "",
                      f"**{payload['course']}**", "",
                      f"**{payload['group']}**", ""]
            lines += [f"- {m}" for m in payload["team"]]
            lines += ["", f"Submission date: {payload['date']}", "", "---", ""]
        elif kind == "toc":
            lines += ["## Table of Contents", "",
                      "*(generated automatically in the DOCX and PDF versions)*", ""]
        elif kind == "h1":
            lines += [f"## {payload}", ""]
        elif kind == "h2":
            lines += [f"### {payload}", ""]
        elif kind == "p":
            lines += [payload, ""]
        elif kind == "bullets":
            lines += [f"- {item}" for item in payload] + [""]
        elif kind == "table":
            lines += [f"**{payload['caption']}**", ""]
            lines.append("| " + " | ".join(payload["headers"]) + " |")
            lines.append("|" + "|".join(["---"] * len(payload["headers"])) + "|")
            for row in payload["rows"]:
                lines.append("| " + " | ".join(row) + " |")
            lines.append("")
        elif kind == "figure":
            rel = Path("figures") / payload["path"].name
            lines += [f"![{payload['caption']}]({rel.as_posix()})", "",
                      f"*{payload['caption']}*", ""]
        elif kind == "pagebreak":
            lines += ["", "---", ""]
    out_path.write_text("\n".join(lines), encoding="utf-8")
    log.info("markdown -> %s", out_path)


# ---------------------------------------------------------------------- docx
def render_docx(blocks: list[Block], out_path: Path) -> None:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.line_spacing = 1.15

    accent = RGBColor(0xB0, 0x3A, 0x2E)

    def _add_field(paragraph, instruction: str) -> None:
        """Insert a Word field (used for the auto table of contents)."""
        run = paragraph.add_run()
        for kind, text in (("begin", None), ("instrText", instruction), ("separate", None)):
            el = OxmlElement(f"w:fld{kind.capitalize()}" if kind != "instrText" else "w:instrText")
            if kind == "begin":
                el = OxmlElement("w:fldChar"); el.set(qn("w:fldCharType"), "begin")
            elif kind == "separate":
                el = OxmlElement("w:fldChar"); el.set(qn("w:fldCharType"), "separate")
            else:
                el.set(qn("xml:space"), "preserve")
                el.text = text
            run._r.append(el)
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run._r.append(end)

    for kind, payload in blocks:
        if kind == "cover":
            for _ in range(4):
                doc.add_paragraph()
            t = doc.add_paragraph()
            t.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = t.add_run(payload["title"])
            run.bold = True
            run.font.size = Pt(26)
            run.font.color.rgb = accent
            for text, size, bold in ((payload["course"], 15, True),
                                     (payload["group"], 13, False)):
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = para.add_run(text)
                r.font.size = Pt(size)
                r.bold = bold
            doc.add_paragraph()
            hdr = doc.add_paragraph()
            hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
            hdr.add_run("Team").bold = True
            for member in payload["team"]:
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.add_run(member).font.size = Pt(12)
            doc.add_paragraph()
            d = doc.add_paragraph()
            d.alignment = WD_ALIGN_PARAGRAPH.CENTER
            d.add_run(f"Submission date: {payload['date']}").font.size = Pt(12)
        elif kind == "toc":
            doc.add_heading("Table of Contents", level=1)
            para = doc.add_paragraph()
            _add_field(para, r'TOC \o "1-2" \h \z \u')
            note = doc.add_paragraph()
            note.add_run("(In Word, right-click the field above and choose "
                         "'Update Field' to populate page numbers.)").italic = True
        elif kind == "h1":
            doc.add_heading(payload, level=1)
        elif kind == "h2":
            doc.add_heading(payload, level=2)
        elif kind == "p":
            doc.add_paragraph(payload)
        elif kind == "bullets":
            for item in payload:
                doc.add_paragraph(item, style="List Bullet")
        elif kind == "table":
            cap = doc.add_paragraph()
            cr = cap.add_run(payload["caption"])
            cr.bold = True
            cr.font.size = Pt(9.5)
            table = doc.add_table(rows=1, cols=len(payload["headers"]))
            table.style = "Light Grid Accent 1"
            for i, head in enumerate(payload["headers"]):
                cell = table.rows[0].cells[i]
                cell.text = str(head)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True
                        r.font.size = Pt(8.5)
            for row in payload["rows"]:
                cells = table.add_row().cells
                for i, val in enumerate(row):
                    cells[i].text = str(val)
                    for p in cells[i].paragraphs:
                        for r in p.runs:
                            r.font.size = Pt(8.5)
            doc.add_paragraph()
        elif kind == "figure":
            doc.add_picture(str(payload["path"]), width=Inches(payload["width"]))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cr = cap.add_run(payload["caption"])
            cr.italic = True
            cr.font.size = Pt(9)
        elif kind == "pagebreak":
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    doc.save(out_path)
    log.info("docx -> %s", out_path)


# ----------------------------------------------------------------------- pdf
def render_pdf(blocks: list[Block], out_path: Path) -> None:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import (Image, ListFlowable, ListItem, PageBreak,
                                    Paragraph, SimpleDocTemplate, Spacer, Table,
                                    TableStyle)
    from PIL import Image as PILImage

    accent = colors.HexColor(ACCENT)
    ss = getSampleStyleSheet()
    body = ParagraphStyle("Body", parent=ss["BodyText"], fontSize=9.5, leading=13.5,
                          alignment=TA_JUSTIFY, spaceAfter=7)
    h1 = ParagraphStyle("H1", parent=ss["Heading1"], fontSize=15, leading=19,
                        textColor=accent, spaceBefore=12, spaceAfter=7)
    h2 = ParagraphStyle("H2", parent=ss["Heading2"], fontSize=12, leading=16,
                        textColor=colors.HexColor("#444444"), spaceBefore=9, spaceAfter=5)
    cap = ParagraphStyle("Cap", parent=body, fontSize=8, alignment=TA_CENTER,
                         textColor=colors.HexColor("#666666"), spaceBefore=3)
    tcap = ParagraphStyle("TCap", parent=body, fontSize=8.5, spaceAfter=3,
                          textColor=colors.HexColor("#333333"))
    cell = ParagraphStyle("Cell", parent=body, fontSize=7, leading=9, alignment=0,
                          spaceAfter=0)
    cellh = ParagraphStyle("CellH", parent=cell, fontName="Helvetica-Bold",
                           textColor=colors.white)
    title_st = ParagraphStyle("T", parent=ss["Title"], fontSize=24, leading=29,
                              textColor=accent)
    center = ParagraphStyle("C", parent=body, alignment=TA_CENTER, fontSize=12)

    def esc(text: str) -> str:
        return (str(text).replace("&", "&amp;").replace("<", "&lt;")
                .replace(">", "&gt;"))

    story: list[Any] = []
    avail_w = LETTER[0] - 2 * inch

    for kind, payload in blocks:
        if kind == "cover":
            story += [Spacer(1, 1.6 * inch),
                      Paragraph(esc(payload["title"]), title_st),
                      Spacer(1, 0.3 * inch),
                      Paragraph(f"<b>{esc(payload['course'])}</b>", center),
                      Paragraph(esc(payload["group"]), center),
                      Spacer(1, 0.35 * inch),
                      Paragraph("<b>Team</b>", center)]
            story += [Paragraph(esc(m), center) for m in payload["team"]]
            story += [Spacer(1, 0.3 * inch),
                      Paragraph(f"Submission date: {esc(payload['date'])}", center)]
        elif kind == "toc":
            story.append(Paragraph("Table of Contents", h1))
            entries = [b[1] for b in blocks if b[0] == "h1"]
            story.append(ListFlowable(
                [ListItem(Paragraph(esc(e), body)) for e in entries],
                bulletType="bullet", leftIndent=14))
        elif kind == "h1":
            story.append(Paragraph(esc(payload), h1))
        elif kind == "h2":
            story.append(Paragraph(esc(payload), h2))
        elif kind == "p":
            story.append(Paragraph(esc(payload), body))
        elif kind == "bullets":
            story.append(ListFlowable(
                [ListItem(Paragraph(esc(i), body), leftIndent=12) for i in payload],
                bulletType="bullet", leftIndent=14, spaceAfter=6))
        elif kind == "table":
            story.append(Paragraph(f"<b>{esc(payload['caption'])}</b>", tcap))
            data = [[Paragraph(esc(h), cellh) for h in payload["headers"]]]
            data += [[Paragraph(esc(c), cell) for c in row] for row in payload["rows"]]
            ncols = len(payload["headers"])
            tbl = Table(data, colWidths=[avail_w / ncols] * ncols, repeatRows=1)
            tbl.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), accent),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#BBBBBB")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1),
                 [colors.white, colors.HexColor("#F6F6F6")]),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 3),
                ("RIGHTPADDING", (0, 0), (-1, -1), 3),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]))
            story += [tbl, Spacer(1, 10)]
        elif kind == "figure":
            path = payload["path"]
            with PILImage.open(path) as im:
                iw, ih = im.size
            w = min(payload["width"] * inch, avail_w)
            h = w * ih / iw
            max_h = 7.0 * inch
            if h > max_h:
                h = max_h
                w = h * iw / ih
            story += [Image(str(path), width=w, height=h),
                      Paragraph(esc(payload["caption"]), cap), Spacer(1, 10)]
        elif kind == "pagebreak":
            story.append(PageBreak())

    def _footer(canvas, doc) -> None:
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(colors.HexColor("#888888"))
        canvas.drawCentredString(LETTER[0] / 2, 0.55 * inch, f"{doc.page}")
        canvas.drawString(inch, 0.55 * inch, "FlameGuard AI - AASD 4014")
        canvas.restoreState()

    SimpleDocTemplate(str(out_path), pagesize=LETTER,
                      leftMargin=inch, rightMargin=inch,
                      topMargin=0.85 * inch, bottomMargin=0.85 * inch,
                      title="FlameGuard AI Final Report").build(
        story, onFirstPage=_footer, onLaterPages=_footer)
    log.info("pdf -> %s", out_path)
